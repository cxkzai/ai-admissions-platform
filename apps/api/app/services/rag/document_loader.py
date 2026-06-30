"""文档加载器 - 支持 Markdown / PDF / DOCX / HTML."""

import re
from dataclasses import dataclass
from pathlib import Path


@dataclass
class LoadedDocument:
    """加载后的文档."""

    content: str
    metadata: dict
    source_path: str
    file_type: str


class DocumentLoader:
    """支持多种格式的文档加载."""

    async def load(self, file_path: str | Path) -> LoadedDocument:
        """根据文件扩展名自动选择加载器."""
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix in {".md", ".markdown"}:
            content = await self._load_markdown(path)
            file_type = "markdown"
        elif suffix == ".txt":
            content = await self._load_text(path)
            file_type = "text"
        elif suffix == ".pdf":
            content = await self._load_pdf(path)
            file_type = "pdf"
        elif suffix in {".docx", ".doc"}:
            content = await self._load_docx(path)
            file_type = "docx"
        elif suffix in {".html", ".htm"}:
            content = await self._load_html(path)
            file_type = "html"
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        # 清洗：去除多余空白
        content = self._clean(content)

        return LoadedDocument(
            content=content,
            metadata={"filename": path.name, "size": len(content)},
            source_path=str(path),
            file_type=file_type,
        )

    async def load_directory(
        self,
        dir_path: str | Path,
        pattern: str = "**/*",
    ) -> list[LoadedDocument]:
        """批量加载目录下所有支持的文档."""
        path = Path(dir_path)
        if not path.exists():
            return []

        docs: list[LoadedDocument] = []
        supported = {".md", ".markdown", ".txt", ".pdf", ".docx", ".doc", ".html", ".htm"}

        for file in path.glob(pattern):
            if file.is_file() and file.suffix.lower() in supported:
                try:
                    doc = await self.load(file)
                    docs.append(doc)
                except Exception as e:
                    print(f"⚠️ Failed to load {file}: {e}")

        return docs

    # ===== 私有方法 =====

    async def _load_markdown(self, path: Path) -> str:
        return path.read_text(encoding="utf-8")

    async def _load_text(self, path: Path) -> str:
        return path.read_text(encoding="utf-8")

    async def _load_pdf(self, path: Path) -> str:
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            return "\n\n".join(page.extract_text() for page in reader.pages)
        except ImportError:
            # 降级到 unstructured
            from unstructured.partition.pdf import partition_pdf

            elements = partition_pdf(str(path))
            return "\n\n".join(str(el) for el in elements)

    async def _load_docx(self, path: Path) -> str:
        from docx import Document

        doc = Document(str(path))
        return "\n\n".join(p.text for p in doc.paragraphs)

    async def _load_html(self, path: Path) -> str:
        from bs4 import BeautifulSoup

        html = path.read_text(encoding="utf-8")
        soup = BeautifulSoup(html, "html.parser")
        # 移除 script/style
        for tag in soup(["script", "style"]):
            tag.decompose()
        return soup.get_text(separator="\n")

    @staticmethod
    def _clean(text: str) -> str:
        """基础清洗：去除多余空白、控制字符、广告标记."""
        # 去除控制字符
        text = re.sub(r"[\x00-\x08\x0b-\x1f\x7f]", "", text)
        # 多空行合并
        text = re.sub(r"\n{3,}", "\n\n", text)
        # 行尾空白
        text = re.sub(r"[ \t]+$", "", text, flags=re.MULTILINE)
        return text.strip()